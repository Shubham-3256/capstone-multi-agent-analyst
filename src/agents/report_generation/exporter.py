"""Exporter compiling Markdown reports into Markdown, HTML, PDF and DOCX documents."""

import os
from pathlib import Path
from typing import Dict, Any

from src.core.logger import get_logger

logger = get_logger(__name__)


class Exporter:
    """Compiles markdown document content into multiple target formats (MD, HTML, PDF, DOCX)."""

    @staticmethod
    def export_markdown(markdown_content: str, target_file: Path) -> Path:
        """Write raw markdown content to target file path.

        Args:
            markdown_content: Raw markdown text.
            target_file: Destination file path.

        Returns:
            Path: Destination file path.
        """
        logger.info(f"Exporter: Exporting Markdown report to: {target_file}")
        target_file.parent.mkdir(parents=True, exist_ok=True)
        with open(target_file, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        return target_file

    @staticmethod
    def export_html(markdown_content: str, target_file: Path) -> Path:
        """Convert markdown text to standard formatted HTML with styling header blocks.

        Args:
            markdown_content: Raw markdown text.
            target_file: Destination file path.

        Returns:
            Path: Destination file path.
        """
        logger.info(f"Exporter: Exporting HTML report to: {target_file}")
        target_file.parent.mkdir(parents=True, exist_ok=True)

        # Simple markdown to HTML parser line-by-line block conversion
        html_lines = []
        for line in markdown_content.splitlines():
            line_strip = line.strip()
            if not line_strip:
                html_lines.append("<br/>")
                continue
            if line_strip.startswith("# "):
                html_lines.append(f"<h1>{line_strip[2:]}</h1>")
            elif line_strip.startswith("## "):
                html_lines.append(f"<h2>{line_strip[3:]}</h2>")
            elif line_strip.startswith("### "):
                html_lines.append(f"<h3>{line_strip[4:]}</h3>")
            elif line_strip.startswith("* "):
                html_lines.append(f"<li>{line_strip[2:]}</li>")
            elif line_strip.startswith("- "):
                html_lines.append(f"<li>{line_strip[2:]}</li>")
            elif line_strip.startswith("**"):
                html_lines.append(f"<p><strong>{line_strip.replace('**', '')}</strong></p>")
            else:
                html_lines.append(f"<p>{line_strip}</p>")

        html_body = "\n".join(html_lines)

        html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Intelligence Report</title>
    <style>
        body {{
            font-family: 'Outfit', 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            color: #1e293b;
            max-width: 800px;
            margin: 40px auto;
            padding: 0 20px;
            background: #f8fafc;
        }}
        h1 {{
            color: #0f172a;
            border-bottom: 2px solid #3b82f6;
            padding-bottom: 8px;
            margin-top: 40px;
        }}
        h2 {{
            color: #1e3a8a;
            margin-top: 30px;
            border-bottom: 1px solid #e2e8f0;
            padding-bottom: 4px;
        }}
        h3 {{
            color: #2563eb;
        }}
        li {{
            margin-bottom: 6px;
        }}
        p {{
            margin-bottom: 16px;
        }}
        strong {{
            color: #0f172a;
        }}
        hr {{
            border: 0;
            border-top: 1px solid #cbd5e1;
            margin: 40px 0;
        }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>
"""
        with open(target_file, "w", encoding="utf-8") as f:
            f.write(html_template)
        return target_file

    @staticmethod
    def export_pdf(markdown_content: str, target_file: Path) -> Path:
        """Compile a styled PDF layout using reportlab library flowables.

        Args:
            markdown_content: Raw markdown text.
            target_file: Destination file path.

        Returns:
            Path: Destination file path.
        """
        logger.info(f"Exporter: Exporting PDF report to: {target_file}")
        target_file.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors

            doc = SimpleDocTemplate(str(target_file), pagesize=letter)
            styles = getSampleStyleSheet()

            # Custom corporate style sheets
            title_style = ParagraphStyle(
                'ReportTitle',
                parent=styles['Heading1'],
                fontName='Helvetica-Bold',
                fontSize=22,
                leading=26,
                textColor=colors.HexColor('#0f172a'),
                spaceAfter=15
            )
            h2_style = ParagraphStyle(
                'ReportH2',
                parent=styles['Heading2'],
                fontName='Helvetica-Bold',
                fontSize=14,
                leading=18,
                textColor=colors.HexColor('#1e3a8a'),
                spaceBefore=12,
                spaceAfter=6
            )
            body_style = ParagraphStyle(
                'ReportBody',
                parent=styles['BodyText'],
                fontName='Helvetica',
                fontSize=10,
                leading=14,
                textColor=colors.HexColor('#334155'),
                spaceAfter=8
            )

            story = []
            for line in markdown_content.splitlines():
                line_strip = line.strip()
                if not line_strip:
                    continue
                if line_strip.startswith("# "):
                    story.append(Paragraph(line_strip[2:], title_style))
                    story.append(Spacer(1, 10))
                elif line_strip.startswith("## ") or line_strip.startswith("### "):
                    # strip tags
                    cleaned = line_strip.replace("##", "").replace("###", "").strip()
                    story.append(Paragraph(cleaned, h2_style))
                    story.append(Spacer(1, 6))
                else:
                    story.append(Paragraph(line_strip, body_style))
                    story.append(Spacer(1, 4))

            doc.build(story)
            logger.info("Exporter: PDF generated successfully via reportlab.")
            
        except Exception as e:
            logger.error(f"Exporter: reportlab compilation failed: {e}. Graceful fallback to text layout.")
            # Graceful fallback: write as clean txt content
            with open(target_file, "w", encoding="utf-8") as f:
                f.write(markdown_content)

        return target_file

    @staticmethod
    def export_docx(markdown_content: str, target_file: Path) -> Path:
        """Compile a Word DOCX file using python-docx library.

        Args:
            markdown_content: Raw markdown text.
            target_file: Destination file path.

        Returns:
            Path: Destination file path.
        """
        logger.info(f"Exporter: Exporting DOCX report to: {target_file}")
        target_file.parent.mkdir(parents=True, exist_ok=True)

        try:
            import docx
            doc = docx.Document()
            
            for line in markdown_content.splitlines():
                line_strip = line.strip()
                if not line_strip:
                    continue
                if line_strip.startswith("# "):
                    doc.add_heading(line_strip[2:], level=1)
                elif line_strip.startswith("## "):
                    doc.add_heading(line_strip[3:], level=2)
                elif line_strip.startswith("### "):
                    doc.add_heading(line_strip[4:], level=3)
                elif line_strip.startswith("* ") or line_strip.startswith("- "):
                    doc.add_paragraph(line_strip[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line_strip)

            doc.save(str(target_file))
            logger.info("Exporter: DOCX generated successfully via python-docx.")

        except Exception as e:
            logger.error(f"Exporter: python-docx compilation failed: {e}. Graceful fallback to text layout.")
            with open(target_file, "w", encoding="utf-8") as f:
                f.write(markdown_content)

        return target_file
